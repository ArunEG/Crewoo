from .models import *
from wkhtmltopdf.views import PDFTemplateResponse
import xlwt
from django.utils.html import strip_tags
from django.http import HttpResponse

def get_or_none(classmodel, **kwargs):
    """
        the function will get the from the Model if exist other
        wise will return none

        @author : Arun Gopi
        @date   : 3/4/2016
    """
    try:
        return classmodel.objects.get(**kwargs)
    except classmodel.DoesNotExist:
        return None

def insertcolx():
    """
            Excel row/colum style for excel
            @Author  : Arun Gopi
            @date    : 3/4/2016
    """

    style = xlwt.XFStyle()
    borders = xlwt.Borders()
    borders.bottom = borders.top = borders.left\
        = borders.right = xlwt.Borders.THIN
    borders.top = xlwt.Borders.THIN
    style.borders = borders
    style.pattern.pattern = 26
    style.pattern.pattern_fore_colour = 0x16
    style.protection.cell_locked = False
    style.protection.formula_hidden = False
    return style


def create_project(data):
    """
        The fnction will create a recored in ProjectInfo table
        @author : Arun
        @date : 9/4/2016

    """
    project = Project()
    project.title = data['title']
    project.description = data['description']
    project.start_date = data['start_date']
    project.end_date = data['end_date']
    project.save()
    return True

def get_project_as_pdf(request, data):
    """ The function will retun project info in pdf format
            @Author  : Arun Gopi
            @date    : 3/4/2016
    """

    filename = 'project_info.pdf'
    response = PDFTemplateResponse(request=request,
                                   template='reports/pdf/project_info.html',
                                   filename=filename,
                                   context=data,
                                   show_content_in_browser=False,
                                   cmd_options={
                                       'encoding': 'utf8', 'quiet': True}
                                   )
    return response

def get_project_as_excel(project):
    """ The function will retun project info in excel format
            @Author  : Arun Gopi
            @date    : 3/4/2016
    """

    row_num, col_num = 0, 0
    work_book = xlwt.Workbook(encoding='utf-8')
    work_sheet = work_book.add_sheet("ProjectInfo")
    insertcol = insertcolx()
    # work_sheet.col(4).hidden = True
    lables = ["Name ", "Overview", "Start date", "Start date"]
    for label in lables:
        work_sheet.write(row_num, col_num, label, insertcol)
        work_sheet.col(col_num).width = 5000
        col_num += 1
    row_num += 1
    work_sheet.write(row_num, 0, str(project.title), insertcol)
    work_sheet.write(
        row_num, 1, str(strip_tags(project.detailed_description)), insertcol)
    work_sheet.write(
        row_num, 2, project.start_date.strftime('%d-%m-%Y %H:%M'), insertcol)
    work_sheet.write(
        row_num, 3, project.end_date.strftime('%d-%m-%Y %H:%M'), insertcol)
    work_sheet.write(
        row_num, 4, project.pk, insertcol)

    response = HttpResponse(content_type='application/ms-excel')
    file_name = "ProjectInfo"+'.xls'
    response['Content-Disposition'] = 'attachment; filename="'+file_name+'"'
    work_book.save(response)
    return response



from docx import Document,shared
from docx.shared import Inches, RGBColor,Pt
import os, subprocess
from docx.oxml.ns import nsdecls
from docx.shared import Inches
from docx.oxml import parse_xml
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_LINE_SPACING



def get_project_as_word(project):
    """ The function will retun project info in word format
            @Author  : Arun Gopi
            @date    : 10/4/2016
    """

    document = Document()

    ''' Setting the top margin '''
    sections = document.sections
    section = sections[0]
    section.top_margin = Inches(0.1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    errFound = False
    mailItenary = False

    dict = {}
    filename = ''
    col1_width = 2
    col2_width = 15
    context_dict = {}
   

    ''' printing important note '''
    font = document.add_paragraph().add_run(project.title).font
    font.size = Pt(10)
    font.underline = True
    font = document.add_paragraph(style='ListBullet').add_run(project.short_description).font
    font.size = Pt(8)
    font = document.add_paragraph(style='ListBullet').add_run(project.detailed_description).font
    font.size = Pt(8)
    font = document.add_paragraph(style='ListBullet').add_run(str(project.start_date)).font
    font.size = Pt(8)
    font = document.add_paragraph(style='ListBullet').add_run(str(project.end_date)).font
    font.size = Pt(8)
    font = document.add_paragraph(style='ListBullet').add_run(project.manager.first_name).font
    font.size = Pt(8)
    font = document.add_paragraph(style='ListBullet').add_run(project.client_name).font
    font.size = Pt(8)
    font = document.add_paragraph(style='ListBullet').add_run(project.createdby.first_name).font
    font.size = Pt(8)
    font = document.add_paragraph(style='ListBullet').add_run(project.modifiedby.first_name).font
    font.size = Pt(8)



    filename = 'project-details'

    filename = filename + '.docx'
    filename = "-".join( filename.split())


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=' + filename
    document.save(response)
    return response


# title
# short_description
# detailed_description
# start_date
# end_date

# manager
# client_name
# createdby
# modifiedby